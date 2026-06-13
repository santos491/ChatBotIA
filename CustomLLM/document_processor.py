import os
from typing import List
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

from config import CHUNK_SIZE, CHUNK_OVERLAP


class DocumentProcessor:
    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""] 
        )
    
    def load_pdf(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        text = ""
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text += f"\n[Página {page_num + 1}]\n{page_text}"
        return text
    
    def load_docx(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text
    
    def load_txt(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def load_document(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self.load_pdf(file_path)
        elif ext == '.docx':
            return self.load_docx(file_path)
        elif ext == '.txt':
            return self.load_txt(file_path)
        else:
            raise ValueError(f"Formato de archivo no soportado por el sistema: {ext}")
    
    def load_from_uploaded_file(self, uploaded_file) -> str:
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_ext == '.pdf':
            reader = PdfReader(uploaded_file)
            text = ""
            for page_num, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += f"\n[Página {page_num + 1}]\n{page_text}"
            return text
        elif file_ext == '.txt':
            return uploaded_file.read().decode('utf-8')
        else:
            raise ValueError(f"Formato de archivo no soportado en la carga web: {file_ext}")
    
    def chunk_text(self, text: str, source: str = "documento_subido") -> List[Document]:
        chunks = self.text_splitter.split_text(text)
        
        documents = []
        for i, chunk in enumerate(chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "source": source,
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            documents.append(doc)
        
        return documents
    
    def process_document(
        self, 
        file_path: str = None, 
        uploaded_file = None,
        source_name: str = None
    ) -> List[Document]:
        if uploaded_file is not None:
            text = self.load_from_uploaded_file(uploaded_file)
            source = source_name or uploaded_file.name
        elif file_path is not None:
            text = self.load_document(file_path)
            source = source_name or os.path.basename(file_path)
        else:
            raise ValueError("Falta información: Se debe proporcionar 'file_path' o 'uploaded_file'.")
        
        return self.chunk_text(text, source)